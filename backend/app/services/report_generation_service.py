"""Build scan-scoped security reports and export them as DOCX or PDF."""

from __future__ import annotations

import html
import io
import json
import re
from collections import Counter
from datetime import datetime, timezone
from typing import Any, Iterable

from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as ReportLabImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


SEVERITIES = ("critical", "high", "medium", "low", "info")
SEVERITY_COLORS = {
    "critical": "#DC2626",
    "high": "#EA580C",
    "medium": "#CA8A04",
    "low": "#2563EB",
    "info": "#64748B",
}
SEVERITY_WEIGHTS = {"critical": 25, "high": 15, "medium": 7, "low": 3}


def _plain(value: Any, fallback: str = "Not provided") -> str:
    """Return renderer-safe text without changing the factual meaning."""
    text = str(value or "").strip()
    if not text:
        return fallback
    return (
        text.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2022", "-")
    )


def _severity(value: Any) -> str:
    normalized = str(value or "info").strip().lower()
    return normalized if normalized in SEVERITIES else "info"


def attack_complexity_from_vector(vector: Any) -> str:
    """Translate the CVSS AC metric without guessing when it is absent."""
    raw = str(vector or "").upper()
    match = re.search(r"(?:^|/)AC:([LHM])(?:/|$)", raw)
    if not match:
        return "Not provided by scanner"
    return {"L": "Low", "M": "Medium", "H": "High"}[match.group(1)]


def _finding_profile(title: str, description: str) -> tuple[str, str, str | None]:
    """Return cautious risk guidance and an optional type-based CWE mapping."""
    text = f"{title} {description}".lower()
    profiles = (
        (
            ("sql injection", "sqli"),
            "Successful exploitation may expose or modify application data and may allow broader database access.",
            "Validate the affected parameter, use parameterized queries, remove unsafe string-built SQL, and retest after remediation.",
            "CWE-89",
        ),
        (
            ("cross-site scripting", " xss", "xss "),
            "Successful exploitation may execute attacker-controlled script in a user's browser and affect sessions or displayed content.",
            "Apply context-aware output encoding, validate untrusted input, use a restrictive Content Security Policy, and retest the affected page.",
            "CWE-79",
        ),
        (
            ("path traversal", "directory traversal"),
            "Successful exploitation may expose files outside the intended application directory.",
            "Canonicalize and allow-list file paths, prevent user-controlled traversal sequences, restrict filesystem permissions, and retest.",
            "CWE-22",
        ),
        (
            ("server-side request forgery", "ssrf"),
            "Successful exploitation may let an attacker make server-side requests to internal or restricted resources.",
            "Allow-list destinations, block private and metadata address ranges, validate redirects, and isolate outbound application traffic.",
            "CWE-918",
        ),
        (
            ("command injection", "remote code execution", " rce"),
            "Successful exploitation may allow unauthorized command execution and compromise the affected service or host.",
            "Remove shell invocation where possible, strictly allow-list inputs, run with least privilege, patch the affected component, and retest.",
            "CWE-78",
        ),
        (
            ("open redirect",),
            "Successful exploitation may redirect users to attacker-controlled locations and support phishing or token leakage scenarios.",
            "Use an allow-list of trusted destinations and reject absolute or scheme-relative user-controlled redirect values.",
            "CWE-601",
        ),
        (
            ("subdomain takeover", "takeover"),
            "A dangling external-service reference may allow an attacker to claim the hostname and serve untrusted content from it.",
            "Remove the dangling DNS record or reclaim and secure the referenced external resource, then verify DNS and HTTP behavior.",
            "CWE-284",
        ),
        (
            ("default credential", "default password"),
            "Default credentials may permit unauthorized access to the exposed service.",
            "Replace all default credentials, enforce unique strong authentication, restrict network access, and review access logs.",
            "CWE-1392",
        ),
        (
            ("denial of service", " dos", "dos "),
            "Successful exploitation may degrade or interrupt availability of the affected service.",
            "Apply the vendor fix or mitigation, restrict unnecessary exposure, add rate and resource controls, and validate resilience.",
            "CWE-400",
        ),
        (
            ("exposed", "disclosure", "sensitive information"),
            "The exposed resource may reveal information that supports unauthorized access or further attack planning.",
            "Remove public exposure where unnecessary, require authorization, minimize returned data, and confirm caches no longer retain it.",
            "CWE-200",
        ),
    )
    for keywords, impact, recommendation, cwe in profiles:
        if any(keyword in text for keyword in keywords):
            return impact, recommendation, cwe
    return (
        "If the scanner match is validated and exploited, it may affect the confidentiality, integrity, or availability of the affected service.",
        "Validate the finding, apply the relevant vendor or configuration remediation, reduce unnecessary exposure, and retest the exact endpoint.",
        None,
    )


def _maturity(counts: dict[str, int], scan_status: str, warning: str) -> dict[str, Any]:
    penalty = sum(counts[name] * weight for name, weight in SEVERITY_WEIGHTS.items())
    score = max(0, 100 - min(100, penalty))
    if warning:
        score = min(score, 85)
    if scan_status.lower() != "completed":
        score = min(score, 70)
    if score >= 90:
        level = "Optimized"
    elif score >= 75:
        level = "Managed"
    elif score >= 50:
        level = "Defined"
    elif score >= 25:
        level = "Developing"
    else:
        level = "Initial"
    return {
        "score": score,
        "level": level,
        "description": (
            "Risk-derived posture indicator based on saved Critical-Low findings. "
            "It is not a compliance certification or a substitute for manual validation."
        ),
    }


def build_scan_report_payload(scan: Any, asset: Any, vulnerabilities: Iterable[Any]) -> dict[str, Any]:
    """Convert one owned scan and its persisted findings into a report payload."""
    rows = [row for row in vulnerabilities if not bool(getattr(row, "is_false_positive", False))]
    rows.sort(
        key=lambda row: (
            SEVERITIES.index(_severity(getattr(row, "severity", None))),
            -(float(getattr(row, "cvss_score", 0) or 0)),
            str(getattr(row, "id", "")),
        )
    )
    counts = Counter(_severity(getattr(row, "severity", None)) for row in rows)
    severity_counts = {name: int(counts.get(name, 0)) for name in SEVERITIES}
    total_findings = sum(severity_counts.values())
    percentages = {
        name: round((severity_counts[name] * 100 / total_findings), 1) if total_findings else 0.0
        for name in SEVERITIES
    }

    actionable_rows = [row for row in rows if _severity(getattr(row, "severity", None)) != "info"]
    findings: list[dict[str, Any]] = []
    for number, row in enumerate(actionable_rows, start=1):
        title = _plain(getattr(row, "title", None), "Untitled scanner finding")
        description = _plain(getattr(row, "description", None), title)
        impact, recommendation, inferred_cwe = _finding_profile(title, description)
        cve = _plain(getattr(row, "cve_id", None), "Not provided by scanner")
        cwe = inferred_cwe or "Not provided by scanner"
        matched_at = _plain(getattr(row, "matched_at", None), "")
        host = _plain(getattr(row, "host", None), "Unknown host")
        port = getattr(row, "port", None)
        affected = matched_at or f"{host}{f':{port}' if port else ''}"
        findings.append({
            "number": number,
            "id": str(getattr(row, "id", "")),
            "title": title,
            "severity": _severity(getattr(row, "severity", None)).capitalize(),
            "cvss_score": float(getattr(row, "cvss_score", 0) or 0) or None,
            "attack_complexity": attack_complexity_from_vector(getattr(row, "cvss_vector", None)),
            "summary": description,
            "potential_impact": impact,
            "recommendation": recommendation,
            "cve": cve,
            "cwe": cwe,
            "cve_cwe": f"CVE: {cve} | CWE: {cwe}",
            "affected_url_ports": affected,
            "source": _plain(getattr(row, "source", None), "scanner"),
        })

    affected_hosts = sorted({
        _plain(getattr(row, "host", None), "")
        for row in rows
        if _plain(getattr(row, "host", None), "")
    })
    affected_ports = sorted({
        int(getattr(row, "port"))
        for row in rows
        if getattr(row, "port", None) is not None
    })
    highest = next((name.capitalize() for name in SEVERITIES if severity_counts[name]), "None")
    warning = _plain(getattr(scan, "error_message", None), "")
    actionable_total = len(actionable_rows)

    observations = [
        {
            "title": "Actionable scanner findings",
            "detail": f"{actionable_total} Critical-Low finding(s) were retained for this scan.",
        },
        {
            "title": "Informational observations",
            "detail": (
                f"{severity_counts['info']} informational observation(s) were recorded and are excluded "
                "from detailed vulnerability sections."
            ),
        },
        {
            "title": "Highest recorded severity",
            "detail": f"The highest severity saved for this scan is {highest}.",
        },
    ]
    if warning:
        observations.append({
            "title": "Scan completion warning",
            "detail": warning,
        })

    generated_at = datetime.now(timezone.utc)
    return {
        "title": "External Attack Surface Security Assessment",
        "generated_at": generated_at,
        "project_objective": (
            "Assess the authorized external target for discoverable security weaknesses and "
            "security-relevant observations, preserve scanner evidence by scan ID, and prioritize "
            "remediation from Critical through Low severity."
        ),
        "scan": {
            "id": str(scan.id),
            "reference_id": _plain(getattr(scan, "reference_id", None), str(scan.id)),
            "status": _plain(getattr(scan, "status", None), "Unknown"),
            "scan_type": _plain(getattr(scan, "scan_type", None), "Unknown"),
            "started_at": getattr(scan, "started_at", None),
            "completed_at": getattr(scan, "completed_at", None),
            "warning": warning or None,
        },
        "target_system": {
            "asset_name": _plain(getattr(asset, "name", None), "Unnamed asset"),
            "asset_type": _plain(getattr(asset, "asset_type", None), "domain"),
            "primary_target": _plain(
                getattr(scan, "target_domain", None) or getattr(asset, "target", None),
                "Not provided",
            ),
            "target_ip": _plain(getattr(scan, "target_ip", None), "Not resolved on scan record"),
            "discovered_assets": int(getattr(scan, "discovered_count", 0) or 0),
            "affected_hosts": affected_hosts,
            "affected_ports": affected_ports,
        },
        "key_observations": observations,
        "maturity": _maturity(severity_counts, str(getattr(scan, "status", "")), warning),
        "severity": {
            "total": total_findings,
            "actionable_total": actionable_total,
            "informational_total": severity_counts["info"],
            "counts": severity_counts,
            "percentages": percentages,
        },
        "findings": findings,
        "report_scope_note": (
            "Detailed observations include Critical, High, Medium, and Low findings only. "
            "Informational detections remain visible in the dashboard percentage but are not "
            "presented as vulnerabilities. Scanner matches should be manually validated before exploitation claims are made."
        ),
    }


def _font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    )
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def build_dashboard_chart(report: dict[str, Any]) -> io.BytesIO:
    """Create a portable dashboard graphic shared by DOCX and PDF exports."""
    image = PILImage.new("RGB", (1200, 420), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(28, bold=True)
    label_font = _font(20, bold=True)
    body_font = _font(18)
    score_font = _font(58, bold=True)
    muted = "#64748B"
    ink = "#0F172A"

    draw.text((36, 24), "Security posture dashboard", fill=ink, font=title_font)
    score = int(report["maturity"]["score"])
    ring_color = "#16A34A" if score >= 75 else "#CA8A04" if score >= 50 else "#DC2626"
    bbox = (55, 90, 325, 360)
    draw.arc(bbox, 0, 360, fill="#E2E8F0", width=30)
    draw.arc(bbox, -90, -90 + round(3.6 * score), fill=ring_color, width=30)
    score_text = str(score)
    score_box = draw.textbbox((0, 0), score_text, font=score_font)
    draw.text(
        (190 - (score_box[2] - score_box[0]) / 2, 166),
        score_text,
        fill=ink,
        font=score_font,
    )
    level = report["maturity"]["level"]
    level_box = draw.textbbox((0, 0), level, font=body_font)
    draw.text(
        (190 - (level_box[2] - level_box[0]) / 2, 238),
        level,
        fill=muted,
        font=body_font,
    )
    draw.text((92, 374), "Maturity score", fill=muted, font=body_font)

    draw.text((390, 92), "Severity distribution", fill=ink, font=label_font)
    counts = report["severity"]["counts"]
    percentages = report["severity"]["percentages"]
    y = 142
    for severity in SEVERITIES:
        color = SEVERITY_COLORS[severity]
        draw.rounded_rectangle((390, y, 535, y + 34), radius=6, fill=color)
        draw.text((407, y + 5), severity.capitalize(), fill="white", font=body_font)
        draw.rounded_rectangle((560, y + 5, 1090, y + 29), radius=8, fill="#E2E8F0")
        width = round(530 * float(percentages[severity]) / 100)
        if width:
            draw.rounded_rectangle((560, y + 5, 560 + max(width, 10), y + 29), radius=8, fill=color)
        label = f"{counts[severity]} finding(s) - {percentages[severity]:.1f}%"
        draw.text((1110, y + 5), label, fill=ink, font=body_font, anchor="ra")
        y += 52

    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def _set_run_font(run: Any, size: float, *, bold: bool = False, color: str = "172033") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_cell_fill(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_geometry(table: Any, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    """Apply compact_reference_guide fixed DXA geometry."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = tc_mar.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _style_table_text(table: Any, header: bool = False) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    _set_run_font(run, 9.5, bold=header and row_index == 0)
        if header and row_index == 0:
            for cell in row.cells:
                _set_cell_fill(cell, "E8EEF5")
            _set_repeat_table_header(row)


def _keep_table_together(table: Any) -> None:
    """Prevent compact finding metadata grids from splitting across pages."""
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index < len(table.rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True


def _add_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    _set_run_font(run, 8.5, color="64748B")


def _add_docx_label_paragraph(document: Document, label: str, value: Any) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, 10.5, bold=True, color="1F4D78")
    value_run = paragraph.add_run(_plain(value))
    _set_run_font(value_run, 10.5)
    return paragraph


def build_docx_report(report: dict[str, Any]) -> bytes:
    """Generate a polished Word report using compact_reference_guide tokens."""
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_table_geometry(header_table, [4680, 4680], indent_dxa=0)
    left = header_table.cell(0, 0).paragraphs[0]
    right = header_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(left.add_run("ASM SECURITY ASSESSMENT"), 8.5, bold=True, color="64748B")
    _set_run_font(right.add_run(report["scan"]["reference_id"]), 8.5, color="64748B")

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(footer_paragraph.add_run("Confidential | Page "), 8.5, color="64748B")
    _add_page_number(footer_paragraph)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(title.add_run(report["title"].upper()), 24, bold=True, color="0B2545")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    _set_run_font(
        subtitle.add_run("Scan-scoped technical findings and remediation report"),
        12,
        color="475569",
    )

    metadata = document.add_table(rows=4, cols=2)
    metadata.style = "Table Grid"
    values = (
        ("Scan ID", report["scan"]["reference_id"]),
        ("Target", report["target_system"]["primary_target"]),
        ("Scan status", report["scan"]["status"]),
        ("Generated", report["generated_at"].strftime("%Y-%m-%d %H:%M UTC")),
    )
    for row, (label, value) in zip(metadata.rows, values):
        row.cells[0].text = label
        row.cells[1].text = _plain(value)
        _set_cell_fill(row.cells[0], "E8EEF5")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    _set_table_geometry(metadata, [1900, 7460])
    _style_table_text(metadata)

    document.add_heading("1. Project objective", level=1)
    document.add_paragraph(report["project_objective"])

    document.add_heading("2. Target system", level=1)
    target = report["target_system"]
    target_table = document.add_table(rows=6, cols=2)
    target_table.style = "Table Grid"
    target_rows = (
        ("Asset", target["asset_name"]),
        ("Primary target", target["primary_target"]),
        ("Target IP", target["target_ip"]),
        ("Discovered assets", target["discovered_assets"]),
        ("Affected hosts", ", ".join(target["affected_hosts"]) or "None recorded"),
        ("Affected ports", ", ".join(map(str, target["affected_ports"])) or "None recorded"),
    )
    for row, (label, value) in zip(target_table.rows, target_rows):
        row.cells[0].text = str(label)
        row.cells[1].text = _plain(value)
        _set_cell_fill(row.cells[0], "E8EEF5")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    _set_table_geometry(target_table, [2700, 6660])
    _style_table_text(target_table)

    document.add_heading("3. Key observations and findings", level=1)
    observation_paragraphs = []
    for observation in report["key_observations"]:
        observation_paragraphs.append(
            _add_docx_label_paragraph(document, observation["title"], observation["detail"])
        )
    for paragraph in observation_paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True

    document.add_heading("4. Security posture dashboard", level=1)
    chart = build_dashboard_chart(report)
    chart_shape = document.add_picture(chart, width=Inches(6.45))
    chart_shape._inline.docPr.set(
        "descr",
        "Security maturity score and Critical, High, Medium, Low, and Info severity percentages.",
    )
    chart_shape._inline.docPr.set("title", "Security posture dashboard")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(report["maturity"]["description"])

    severity_table = document.add_table(rows=1, cols=3)
    severity_table.style = "Table Grid"
    severity_table.rows[0].cells[0].text = "Severity"
    severity_table.rows[0].cells[1].text = "Count"
    severity_table.rows[0].cells[2].text = "Percentage"
    for severity in SEVERITIES:
        cells = severity_table.add_row().cells
        cells[0].text = severity.capitalize()
        cells[1].text = str(report["severity"]["counts"][severity])
        cells[2].text = f"{report['severity']['percentages'][severity]:.1f}%"
    _set_table_geometry(severity_table, [4260, 2200, 2900])
    _style_table_text(severity_table, header=True)
    document.add_paragraph(report["report_scope_note"])

    document.add_page_break()
    document.add_heading("5. Detailed observation details", level=1)
    if not report["findings"]:
        document.add_paragraph(
            "No Critical, High, Medium, or Low findings were saved for this scan. "
            "Informational observations are intentionally excluded from this section."
        )
    for finding in report["findings"]:
        document.add_heading(
            f"Finding {finding['number']}: {finding['title']}",
            level=2,
        )
        detail_table = document.add_table(rows=5, cols=2)
        detail_table.style = "Table Grid"
        detail_rows = (
            ("Finding ID", finding["id"]),
            ("Severity", finding["severity"]),
            ("Attack complexity", finding["attack_complexity"]),
            ("CVE / CWE", finding["cve_cwe"]),
            ("Affected URL and ports", finding["affected_url_ports"]),
        )
        for row, (label, value) in zip(detail_table.rows, detail_rows):
            row.cells[0].text = label
            row.cells[1].text = _plain(value)
            _set_cell_fill(row.cells[0], "E8EEF5")
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
        _set_table_geometry(detail_table, [2700, 6660])
        _style_table_text(detail_table)
        _keep_table_together(detail_table)
        _add_docx_label_paragraph(document, "Finding summary", finding["summary"])
        _add_docx_label_paragraph(document, "Potential impact", finding["potential_impact"])
        _add_docx_label_paragraph(document, "Recommendation", finding["recommendation"])

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_text(value: Any) -> str:
    return html.escape(_plain(value)).replace("\n", "<br/>")


def build_pdf_report(report: dict[str, Any]) -> bytes:
    """Generate a polished PDF report with the same structure as the Word export."""
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.82 * inch,
        bottomMargin=0.72 * inch,
        title=report["title"],
        author="ASM Platform",
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "subtitle": ParagraphStyle(
            "ReportSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor("#475569"),
            spaceAfter=16,
        ),
        "h1": ParagraphStyle(
            "ReportH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "ReportH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "ReportBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#172033"),
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "ReportSmall",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#475569"),
        ),
        "label": ParagraphStyle(
            "ReportLabel",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=2,
        ),
    }

    table_style = TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#CBD5E1")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ])

    def paragraph(value: Any, style: str = "body") -> Paragraph:
        return Paragraph(_pdf_text(value), styles[style])

    def label_value(label: str, value: Any) -> Paragraph:
        return Paragraph(
            f"<b><font color='#1F4D78'>{_pdf_text(label)}:</font></b> {_pdf_text(value)}",
            styles["body"],
        )

    story: list[Any] = [
        Paragraph(_pdf_text(report["title"].upper()), styles["title"]),
        Paragraph("Scan-scoped technical findings and remediation report", styles["subtitle"]),
    ]
    metadata_data = [
        [paragraph("Scan ID", "label"), paragraph(report["scan"]["reference_id"], "small")],
        [paragraph("Target", "label"), paragraph(report["target_system"]["primary_target"], "small")],
        [paragraph("Scan status", "label"), paragraph(report["scan"]["status"], "small")],
        [paragraph("Generated", "label"), paragraph(report["generated_at"].strftime("%Y-%m-%d %H:%M UTC"), "small")],
    ]
    metadata = Table(metadata_data, colWidths=[1.35 * inch, 5.15 * inch], repeatRows=0)
    metadata.setStyle(table_style)
    story.extend([metadata, Spacer(1, 8)])

    story.extend([
        Paragraph("1. Project objective", styles["h1"]),
        paragraph(report["project_objective"]),
        Paragraph("2. Target system", styles["h1"]),
    ])
    target = report["target_system"]
    target_data = [
        [paragraph("Asset", "label"), paragraph(target["asset_name"], "small")],
        [paragraph("Primary target", "label"), paragraph(target["primary_target"], "small")],
        [paragraph("Target IP", "label"), paragraph(target["target_ip"], "small")],
        [paragraph("Discovered assets", "label"), paragraph(target["discovered_assets"], "small")],
        [paragraph("Affected hosts", "label"), paragraph(", ".join(target["affected_hosts"]) or "None recorded", "small")],
        [paragraph("Affected ports", "label"), paragraph(", ".join(map(str, target["affected_ports"])) or "None recorded", "small")],
    ]
    target_table = Table(target_data, colWidths=[1.75 * inch, 4.75 * inch])
    target_table.setStyle(table_style)
    story.extend([target_table, Paragraph("3. Key observations and findings", styles["h1"])])
    story.extend(label_value(item["title"], item["detail"]) for item in report["key_observations"])

    story.append(Paragraph("4. Security posture dashboard", styles["h1"]))
    chart = build_dashboard_chart(report)
    story.append(ReportLabImage(chart, width=6.45 * inch, height=2.26 * inch))
    story.append(paragraph(report["maturity"]["description"], "small"))
    severity_data = [[paragraph("Severity", "label"), paragraph("Count", "label"), paragraph("Percentage", "label")]]
    for severity in SEVERITIES:
        severity_data.append([
            paragraph(severity.capitalize(), "small"),
            paragraph(report["severity"]["counts"][severity], "small"),
            paragraph(f"{report['severity']['percentages'][severity]:.1f}%", "small"),
        ])
    severity_table = Table(severity_data, colWidths=[3 * inch, 1.5 * inch, 2 * inch], repeatRows=1)
    severity_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#CBD5E1")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend([severity_table, Spacer(1, 5), paragraph(report["report_scope_note"], "small")])

    story.extend([PageBreak(), Paragraph("5. Detailed observation details", styles["h1"])])
    if not report["findings"]:
        story.append(paragraph(
            "No Critical, High, Medium, or Low findings were saved for this scan. "
            "Informational observations are intentionally excluded from this section."
        ))
    for finding in report["findings"]:
        story.append(Paragraph(
            _pdf_text(f"Finding {finding['number']}: {finding['title']}"),
            styles["h2"],
        ))
        details = [
            [paragraph("Finding ID", "label"), paragraph(finding["id"], "small")],
            [paragraph("Severity", "label"), paragraph(finding["severity"], "small")],
            [paragraph("Attack complexity", "label"), paragraph(finding["attack_complexity"], "small")],
            [paragraph("CVE / CWE", "label"), paragraph(finding["cve_cwe"], "small")],
            [paragraph("Affected URL and ports", "label"), paragraph(finding["affected_url_ports"], "small")],
        ]
        detail_table = Table(details, colWidths=[1.75 * inch, 4.75 * inch])
        detail_table.setStyle(table_style)
        story.extend([
            detail_table,
            Spacer(1, 5),
            label_value("Finding summary", finding["summary"]),
            label_value("Potential impact", finding["potential_impact"]),
            label_value("Recommendation", finding["recommendation"]),
            Spacer(1, 8),
        ])

    def draw_page(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#CBD5E1"))
        canvas.setLineWidth(0.5)
        canvas.line(inch, letter[1] - 0.52 * inch, letter[0] - inch, letter[1] - 0.52 * inch)
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawString(inch, letter[1] - 0.42 * inch, "ASM SECURITY ASSESSMENT")
        canvas.setFont("Helvetica", 7.5)
        canvas.drawRightString(letter[0] - inch, letter[1] - 0.42 * inch, report["scan"]["reference_id"])
        canvas.line(inch, 0.5 * inch, letter[0] - inch, 0.5 * inch)
        canvas.drawRightString(letter[0] - inch, 0.35 * inch, f"Confidential | Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return output.getvalue()


def report_to_json(report: dict[str, Any]) -> str:
    """Stable JSON helper used only for diagnostics and tests."""
    return json.dumps(report, default=str, indent=2)
